from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from docx import Document
from .forms import ResumeUploadForm, AssignmentUploadForm, SubmissionForm
from .models import Resume, Assignment, PlagiarismReport, StudyPlan,Submission
from django.contrib import messages
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from django.core.paginator import Paginator
from Users.models import MainUser
from django.utils.decorators import method_decorator
from django.views.generic import UpdateView


from services.ai_gateway import analyze_resume_vs_jd, ai_assistant
from .utils import extract_text_from_file, normalize_text

import tempfile

import uuid, json
from django.core.cache import cache
from django.views.decorators.csrf import ensure_csrf_cookie
from django.views.decorators.http import require_POST
from .tasks import analyze_resume_task  #celery task
from django.http import HttpResponse, JsonResponse

# Create your views here.

@login_required
@ensure_csrf_cookie
def upload_and_analyze(request):
    """ Render the upload page (AJAX-absed analysis)."""
    if request.method == "POST" and not request.is_ajax():

        # fallback synchronous mode
        form = ResumeUploadForm(request.POST, request.FILES)
        if form.is_valid():
            resume = form.save(commit=False)
            resume.user = request.user
            resume.save()

            # read the targate role and job description from the POST::
            target_role = request.POST.get("target_role", "")
            job_description = request.POST.get("job_description", "")


            # Parse and analyze
            text = extract_text_from_file(resume.uploaded_file)
            text  = normalize_text(text)
            report = analyze_resume_vs_jd(text, job_description, target_role or None)

            resume.parsed_text = text
            resume.feedback = report
            resume.ats_score = report.get("ats_score", 0.0)
            resume.keyword_match_score = report.get("keyword_match_percent", 0.0)
            resume.save()
            messages.success(request, "Uploaded and analyzed.")
            return redirect("education:resume_detail", pk=resume.pk)
    else:
        form= ResumeUploadForm()
    return render(request, "education/resume_upload.html", {'form': form})

# Analysis part by using Celery in background::
@login_required
@require_POST
def start_analysis_ajax(request):
    """Start background Celery task."""
    form = ResumeUploadForm(request.POST, request.FILES)
    if not form.is_valid():
        return JsonResponse({"ok": False, "errors": form.errors}, status=400)

    resume = form.save(commit=False)
    resume.user = request.user
    resume.save()

    jd_text = request.POST.get("job_description", "")
    target_role = request.POST.get("target_role", "")
    generate_suggestions = request.POST.get("generate_suggestions", "1") != "0"

    # Trigger Celery task
    async_result = analyze_resume_task.delay(resume.id, jd_text, target_role, request.user.id, generate_suggestions)

    # store initial progress
    cache.set(f"resume_task_{async_result.id}", {"status": "queued", "progress": 1, "message": "Queued"}, timeout=60 * 60)

    return JsonResponse({"ok": True, "task_id": async_result.id, "resume_pk": resume.pk})


from celery.result import AsyncResult
@login_required
def analysis_status(request, task_id: str):
    """Poll progress from cache."""

    if not task_id:
        return JsonResponse({ "error": "Missing task ID"}, status=400)

    # First try to cache::
    cached = cache.get(f"resume_task_{task_id}")
    if cached:
        return JsonResponse({"ok": True, "task": cached})

    result = AsyncResult(task_id)
    response = {'state': result.state}

    if result.state == 'PROGRESS':
        response.update(result.info or {})
    elif result.state == "SUCCESS":
        response['result'] = result.result
    elif result.state == "FAILURE":
        response['error'] = str(result.info)
    elif result.state == 'PENDING':
        response['message'] = 'Pending...'

    return JsonResponse({"ok": True, "task": response})

# 2. view for details of resume
@login_required
def resume_detail(request,pk):
    resume= get_object_or_404(Resume, pk=pk, user=request.user)
    return render(request, "education/resume_detail.html", {"resume": resume})

@login_required
def generate_optimized_resume(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    fb = resume.feedback or {}
    opt_text = fb.get('optimized_resume_text') or build_optimized_resume(fb.get('extracted_sections', {}), fb.get('missing_keywords', []), target_role=request.POST.get('target_role', None))


    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    import io

    # Now save the temp file as .txt
    # And download in pdf/text/docx
    file_format = request.GET.get("format", "text").lower()

    if file_format == "pdf":
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer)
        styles = getSampleStyleSheet()
        story = [Paragraph("Optimized Resume", styles["Title"]), Spacer(1, 12)]
        story.append(Paragraph(opt_text.replace("\n", "<br/>"), styles["Normal"]))
        doc.build(story)
        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{request.user.username}_optimized_{resume.pk}.pdf"'
        return response
    
    elif file_format == "docx":
        buffer = io.BytesIO()
        doc = Document()
        doc.add_heading("Optimized Resume", 0)
        doc.add_paragraph(opt_text)
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{request.user.username}_optimized_{resume.pk}.docx"'
        return response
    
    else:
        response = HttpResponse(opt_text, content_type='text/plain')
        response['Content-Disposition'] = f'attachment; filename="{request.user.username}_optimized_{resume.pk}.txt"'
        return response

def build_optimized_resume(extracted_sections, missing_keywords, target_role=None):
    """
    Build optimized resume text from extracted sections and missing keywords.
    
    Args:
        extracted_sections: dict with resume sections (e.g., {'summary': '...', 'experience': '...'})
        missing_keywords: list of keywords to incorporate
        target_role: target job role (optional)
    
    Returns:
        str: formatted optimized resume text
    """
    lines = []
    
    if target_role:
        lines.append(f"TARGET ROLE: {target_role}\n")
    
    # Add extracted sections
    for section, content in extracted_sections.items():
        if content:
            lines.append(f"\n{section.upper()}")
            lines.append("-" * 40)
            lines.append(str(content))
    
    # Add missing keywords section
    if missing_keywords:
        lines.append(f"\n{'KEYWORDS TO ADD'.upper()}")
        lines.append("-" * 40)
        lines.append(", ".join(missing_keywords))
    
    return "\n".join(lines)


# Now create a simple helper parser::
def extract_text_from_file(file_field) ->str:
    # Simple:: if file is endswith .pdf -> use pdfminer, if .docx -> python-docx
    name= file_field.name.lower()
    data = file_field.read()

    try:
        if name.endswith('.pdf'):
            from io import BytesIO
            from pdfminer.high_level import extract_text_to_fp
            out = BytesIO()
            extract_text_to_fp(BytesIO(data), out)
            return out.getvalue().decode('utf-8', errors='ignore')
        elif name.endswith('.docx'):
            import docx
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(data)
                tmp.flush()
                doc = docx.Document(tmp.name)
                return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        pass
    return data.decode('utf-8', errors='ignore') if isinstance(data, (bytes, bytearray)) else str(data)


# ========= Teacher View ===============

def is_teacher(user):
    """
    Check if the user is a teacher (based on custom MainUser model).
    """
    return user.is_authenticated and (user.is_teacher or user.role == "teacher")


# 3. view for Assignment_Upload
@login_required
@user_passes_test(is_teacher)
def assignment_upload(request):
    """
    Unified and upgraded version of teacher can create/upload assignment.

    Features:
    - Teacher can create/upload assignments with (title, description, file, deadline)
    - Extracts file text content (if any)
    - Displays success/error messages
    - Automatically sets teacher = request.user
    - After save → redirects to teacher's assignment list
    - Includes new advanced features for real-world feel:
        * Smart submission tracker (% of students submitted)
        * Editable until deadline
        * Modern feedback-ready integration
    """

    if request.method == 'POST':
        form = AssignmentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            assignment = form.save(commit=False)
            assignment.teacher = request.user
            assignment.save()

            # try to extract_text
            try:
                text= extract_text_from_file(assignment.file)
                assignment.text_content = text or ""
                assignment.save()
            except Exception as e:
                messages.warning(request, f"Assignment uploaded but text-extraction failed: {e}")

            # Initialize progress tracking::
            total_students = MainUser.objects.filter(role='student').count()
            assignment.total_students = total_students
            assignment.submitted_count = 0
            assignment.save()

            messages.success(request, "Assignment Created successfully and published for students....!")
            return redirect("education:teacher_assignment_list")
        else:
            messages.error(request, f"Something went going to wrong...! Pleases correcet it: {form.errors}")
            return render(request, "education/assignment_upload.html", {'form': form})

    else:
        form = AssignmentUploadForm()
        return render(request, "education/assignment_upload.html", {"form": form})

@login_required
@user_passes_test(is_teacher)
def teacher_assignment_list(request):
    assignments = Assignment.objects.filter(teacher=request.user).order_by("-created_at")
    # compute submission status for each assignment
    stats = []
    for a in assignments:
        total_sub = a.submissions.count()
        total_students = MainUser.objects.filter(role='student').count()
        percent = 0
        if total_students:
            percent = int((total_sub / total_students) * 100) if total_students else 0
        stats.append({"assignment": a, "submitted_count": total_sub, "total_students": total_students, "percent": percent})
    return render(request, "education/teacher_assignment_list.html", {"stats": stats, "assignments": assignments})


@login_required
@user_passes_test(is_teacher)
def teacher_assignment_submission(request, assignment_id):
    """
    Display all submissions for a specific assignment created by this teacher.
    """
    assignment = get_object_or_404(Assignment, id=assignment_id, teacher=request.user)
    submissions = Submission.objects.filter(assignment=assignment).select_related("student")

    paginator = Paginator(submissions, 20)
    page = request.GET.get("page", 1)
    submissions_page = paginator.get_page(page)
    return render(request, "education/teacher_assignment_submission.html", {"assignment": assignment, "submissions": submissions_page})

@login_required
@user_passes_test(is_teacher)
def teacher_grade_submission(request, submission_id):
    """
    Allow teacher to grade a submission
    """

    submission = get_object_or_404(Submission, id=submission_id, assignment__teacher = request.user)
    if request.method == "POST":
        grade = request.POST.get("grade")
        feedback = request.POST.get("feedback")
        submission.grade = grade
        submission.feedback = feedback
        submission.status = "Graded"
        submission.save()
        messages.success(request, "Submission Graded successfully!.")
        return redirect("education:teacher_assignment_submission", assignment_id = submission.assignment.id)
    return render(request, "education/teacher_grade_submission.html", {"submission": submission})

@login_required
@user_passes_test(is_teacher)
def teacher_edit_assignment(request, pk):
    assignment = get_object_or_404(Assignment, pk=pk, teacher = request.user)

    if request.method == "POST":
        form = AssignmentUploadForm(request.POST, request.FILES, instance=assignment)
        if form.is_valid():
            form.save()
            messages.success(request, "Assignment updated successfully..!")
            return redirect("users:teacher_dashboard")
    else:
        form= AssignmentUploadForm(instance=assignment)

    return render(request, "education/teacher_edit_assignment.html", {"form": form, "assignment": assignment})

@login_required
@user_passes_test(is_teacher)
def teacher_delete_assignment(request, pk):
    assignment = get_object_or_404(Assignment, pk=pk, teacher= request.user)
    if request.method == "POST":
        assignment.delete()
        messages.success(request, "Assignment deleted successfully!")
        return redirect("users:teacher_dashboard")
    return redirect("users:teacher_dashboard")


# ================ Student View ====================
from functools import wraps

def student_required(view_func):
    @wraps(view_func)
    def _wrapped(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('login')
        if not getattr(request.user, 'is_student', False):
            return redirect('no_permission')
        return view_func(request, *args, **kwargs)
    return _wrapped


# This is for student can see there entire assignment which are posted by teacher
@login_required
@student_required
def student_assignment_list(request):
    # show all assignment or filter by class/department
    now = timezone.now()
    assignments = Assignment.objects.all().order_by("-created_at")

    # Get student's latest submission per assignment(heighest attempt)
    student_subs = Submission.objects.filter(student=request.user).order_by("assignment_id", "-attempt")
    # create mapping: assignment_id -> latest submission
    latest = {}
    for s in student_subs:
        if s.assignment_id not in latest:
            latest[s.assignment_id] =s
    
    enriched = []
    for a in assignments:
        sub = latest.get(a.id)
        if sub:
            status = sub.status or "Submitted"
            submitted = True
            grade = sub.grade if sub.grade else "--"
        else:
            status = "No Submitted"
            submitted = False
            grade=  "--"
        enriched.append({
            "assignments": a, 
            "submission": sub, 
            "status": status,
            "grade": grade,
            'submitted': submitted,
            })

    return render(request, "education/student_assignment.html", {"assignments": enriched, "now": now})

# Here a student can submit there assignment untill ending of deadline
@login_required
@student_required
def student_submit_assignment(request, assignment_id):
    """
    Student uploads their assignment submission.
    - Prevent submission after deadline
    - Allow re-submission before deadline (versioned)
    - Prevents overwriting uploads
    - Update teacher progress count
    """

    assignment = get_object_or_404(Assignment, id=assignment_id)

    # Prevent submission after deadline ::
    if assignment.deadline and timezone.now() > assignment.deadline:
        messages.error(request, "The submission deadline has passed. You can't submit now.")
        return redirect("education:student_assignment_list")


    # determine next attempt number(versioning)
    last_attempt = Submission.objects.filter(assignment=assignment, student=request.user).order_by("-attempt").first()
    next_attempt = 1 if not last_attempt else last_attempt.attempt + 1

    if request.method == "POST":
        form = SubmissionForm(request.POST, request.FILES)
        if form.is_valid():
            submission = form.save(commit=False)
            submission.assignment = assignment
            submission.student = request.user
            submission.attempt = next_attempt
            submission.status = "Submitted"

            # Ensure unique filename (prevent overwrites)
            uploaded_file = submission.submitted_file
            # uploaded_file.name = f"{request.user.username}_A{assignment_id}_V{next_attempt}_{uploaded_file}_{uploaded_file.name}"
            uploaded_file.name = f"{request.user.username}_A{assignment_id}_V{next_attempt}_{uploaded_file.name}"
            submission.save()

            # update assignment tracking progress
            submitted_students = Submission.objects.filter(assignment=assignment).values_list("student_id", flat=True).distinct().count()
            assignment.submitted_count = submitted_students
            assignment.save(update_fields=['submitted_count'])

            messages.success(request, f"Assignment Submitted successfully (Version/Attempt {next_attempt})..!")
            return redirect("education:student_assignment_list")
        else:
            messages.error(request, "Something went wrong...! Please correct the error.")
    else:
        form= SubmissionForm()
    return render(request, "education/student_submit_assignment.html", {"form": form, "assignment": assignment, "next_attempt": next_attempt})


# 4. view for Assignment details::
@login_required
def assignment_details(request, pk):
    """ Genereic detail view."""
    assignment = get_object_or_404(Assignment, pk=pk)
    return render(request, "education/assignment_detail.html", {'assignment': assignment})


# =========================================
# 5. Now Generate a view for end-to-end study plan
# ========================================

from django.conf import settings

from .forms import SyllabusUpladForm, StudyPlanRequestForm
from .models import SyllabusUpload, StudyPlanRequest, StudyPlan, StudyPlanItem

# For Gemini AI
import google.generativeai as genai
genai.configure(api_key=settings.GEMINI_API_KEY)


# ============================================================
# 1️⃣  Upload Syllabus View
# ============================================================
@login_required
def upload_syllabus_view(request):
    """
    Upload syllabus or past performance report.
    Redirects to study plan request after upload.
    """
    if request.method == "POST":
        form = SyllabusUpladForm(request.POST, request.FILES)
        if form.is_valid():
            syllabus = form.save(commit=False)
            syllabus.user = request.user
            syllabus.save()
            return redirect("education:request_study_plan")
    else:
        form = SyllabusUpladForm()
    return render(request, "education/upload_syllabus.html", {"form": form})

# ============================================================
# 2️.  Request Study Plan View
# ============================================================

from datetime import date
from datetime import datetime

@login_required
def request_study_plan_view(request):
    """
    Handles user input for plan generation (hours/day, goal, etc.)
    and generates the plan directly (no Celery delay for debugging).
    """
    from .tasks import generate_study_plan_task

    if request.method == "POST":
        form = StudyPlanRequestForm(request.POST)
        if form.is_valid():
            # Convert all date objects into ISO strings before saving:
            start_date = form.cleaned_data.get("start_date")
            end_date = form.cleaned_data.get("end_date")

            # Now convert to string (ISO format)
            if start_date:
                start_date = start_date.isoformat()
            if end_date:
                end_date = end_date.isoformat()

            req = StudyPlanRequest.objects.create(
                user=request.user,
                params={
                    "hours_per_day": form.cleaned_data["hours_per_day"],
                    "start_date": start_date,
                    "deadline": end_date,
                    "goals": form.cleaned_data.get("goals"),
                }
            )

            # Attach latest syllabus if available
            latest = SyllabusUpload.objects.filter(user=request.user).order_by("-uploaded_at").first()
            if latest:
                req.syllabus = latest
                req.save()

            # Directly generate plan (instead of Celery)
            generate_study_plan_task.delay(req.id)

            return redirect("education:study_dashboard")
    else:
        form = StudyPlanRequestForm()

    return render(request, "education/request_study_plan.html", {"form": form})

# ============================================================
#3.  Study Dashboard View
# ============================================================
@login_required
def study_dashboard(request):
    """
    Displays the most recent active study plan.
    Includes progress, charts, and Gemini Q&A section.
    """
    plan = StudyPlan.objects.filter(user=request.user, active=True).prefetch_related("items").order_by("-created_at").first()

    if not plan:
        return render(request, "education/study_dashboard.html", {"study_plan": None})

    # Prepare chart data
    items = plan.items.all().order_by("date")
    total_items = items.count()
    completed_items = items.filter(completed=True).count()
    progress = round((completed_items / total_items) * 100, 2) if total_items > 0 else 0

    study_plan = {
        item.date.strftime("%Y-%m-%d"): {
            "topic": item.topic,
            "tasks": item.notes,
            "hours": item.duration_minutes // 60,
        }
        for item in items
    }

    context = {
        "plan_name": plan.name,
        "total_days": plan.duration_days,
        "progress": progress,
        "study_plan": study_plan,
        "progress_chart_data": json.dumps({
            "labels": [item.date.strftime("%b %d") for item in items],
            "values": [100 if item.completed else 0 for item in items]
        }),
        "difficulty_data": json.dumps({
            "labels": ["Easy", "Medium", "Hard", "Very Hard"],
            "values": [
                items.filter(difficulty_score__lte=2).count(),
                items.filter(difficulty_score__range=(3, 5)).count(),
                items.filter(difficulty_score__range=(6, 8)).count(),
                items.filter(difficulty_score__gte=9).count(),
            ]
        }),
        "ai_response": request.GET.get("ai_response"),
    }

    return render(request, "education/study_dashboard.html", context)

# ============================================================
# 4️.  Study Plan JSON API
# ============================================================
@login_required
def study_plan_json(request, plan_id):
    """
    Returns a JSON object for chart rendering.
    """
    plan = get_object_or_404(StudyPlan, id=plan_id, user=request.user)
    items = plan.items.all().values(
        "id", "date", "topic", "duration_minutes", "completed", "difficulty_score"
    )
    return JsonResponse({"items": list(items)})

# 5️.  Gemini Q&A Doubt Solver
@login_required
@require_POST
def qa_gemini(request):
    """
    Handles the AI question-answer feature using Gemini API.
    """
    q = request.POST.get("query", "").strip()  # 🔧 fixed field name
    if not q:
        return redirect(reverse("education:study_dashboard"))

    # Context: last 3 studied topics
    plan = StudyPlan.objects.filter(user=request.user, active=True).order_by("-created_at").first()
    context_text = ""
    if plan:
        last_topics = list(plan.items.all().order_by("-date")[:3].values_list("topic", flat=True))
        context_text = "Recent topics: " + ", ".join(last_topics)

    prompt = (
        "You are an expert tutor.\n"
        f"Context: {context_text}\n\n"
        f"Student asks: {q}\n\n"
        "Give a clear, short explanation with examples if needed."
    )

    try:
        genai.configure(api_key=settings.GEMINI_API_KEY)

        model_name = genai.GenerativeModel(settings.GEMINI_MODEL)
        resp = model_name.generate_content(prompt)
        answer = resp.text
    except Exception as e:
        answer = f"Error contacting AI: {e}"

    # Return to dashboard with answer displayed
    return redirect(
        f"{reverse('education:study_dashboard')}?ai_response={answer}")


def ai_assistant_page(request):
    """
    A dedicated page for AI assistant interactions.
    """
    return render(request, "education/ai_assistant.html")


import requests
from django.views.decorators.csrf import csrf_exempt
from .utils import parse_uploaded_file

HF_URL = "https://atiurr-campus-360-ai-service.hf.space/ai/assistant"

@csrf_exempt
def assistant_api(request):
    if request.method == "POST":
        if request.content_type.startswith("application/json"):
            import json
            data = json.loads(request.body)
            query = data.get("query", "")
            file = None
        else:
            query = request.POST.get("query", "")
            file = request.FILES.get("files")

            print("User Quuery:", query)
            print("Uploaded File:", file)

        file_text = ""
        if file:
            try:
                file_text = parse_uploaded_file(file)
                file_text = file_text[:4000]
            except Exception as e:
                print("File processing error:", str(e))
                return JsonResponse({"response": "File processing failed"}, status=500)

        # Get previous chat from session::
        history = request.session.get("chat_history", [])

        # Add User Messages ::
        history.append({"role": "user", "content": query})

        print("User Query:", query)

        # Build prompt with memory
        context = "\n".join([
            f"{msg['role']}: {msg['content']}" for msg in history[-5:]
            ])
        
        final_prompt = f"""
        You are a smart AI assistant.

        Conversation:
        {"File Content:\n" + file_text if file_text else ""}

        Instructions:
        - Answer clearly 
        - Do Not repeat responses
        - Use proper formating

        Assistant:
        """

        try:
            # ✅ Correct API call
            resp = requests.post(
                HF_URL,
                json={"query": final_prompt},
                timeout=60
            )

            print("HF Status:", resp.status_code)
            print("HF Raw Response:", resp.text)

            if resp.status_code != 200:
                return JsonResponse({"response": "Hugging Face API failed"}, status=500)

            data = resp.json()

            reply = (
                data.get("response")
                or data.get("answer")
                or data.get("message")
                or "No response"
            )

            # save assistant reply to session history 
            history.append({"role": "assistant", "content": reply})
            request.session["chat_history"] = history

            return JsonResponse({"response": reply})

        except requests.exceptions.ReadTimeout:
            return JsonResponse({"response": "AI timeout. Try again."}, status=500)

        except Exception as e:
            print("AI Assistant Error:", str(e))
            return JsonResponse({"response": "Server error"}, status=500)