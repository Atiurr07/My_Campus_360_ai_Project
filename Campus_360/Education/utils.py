import os
from django.core.files.uploadedfile import UploadedFile
from django.db.models.fields.files import FieldFile

import PyPDF2
import docx

# File Text Extraction::
def extract_text_from_file(file_or_path) -> str:

    """
    Supports:
    - UploadedFile (request.Files)
    - FieldFile(model FileField)
    - str / Path (Celery-safe)
    """

    if isinstance(file_or_path, UploadedFile):
        ext = os.path.splitext(file_or_path.name)[1].lower()
        return _extract_uploaded(file_or_path, ext)
    
    if isinstance(file_or_path, FieldFile):
        return extract_text_from_file(file_or_path.path)

    if isinstance(file_or_path, (str, os.PathLike)):
        path = str(file_or_path)
        ext = os.path.splitext(path)[1].lower()
        return _extract_from_path(path, ext)
    
    raise TypeError(f"Unsupported input: {type(file_or_path)}")

def _extract_uploaded(file, ext):
    if ext == ".pdf":
        reader = PyPDF2.PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    
    if ext in [".docx", ".doc"]:
        document = docx.Document(file)
        return "\n".join(p.text for p in document.paragraphs)
    
    if ext == ".txt":
        return file.read().decode("utf-8")

    raise ValueError("Unsupported file format")

def _extract_from_path(path, ext):
    if not os.path.exists(path):
        raise FileNotFoundError(path)

    if ext == ".pdf":
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

    if ext in [".docx", "doc"]:
        document = docx.Document(path)
        return "\n".join(p.text for p in document.paragraphs)

    if ext == ".txt":
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    
    raise ValueError(f"Unsupported file format: {ext}")

def normalize_text(text: str) -> str:
    return (
        text.replace("\x00", "")
        .replace("\r", "")
        .strip()
    )

from datetime import datetime, date, timedelta

def flatten_ai_plan_to_items(ai_json: dict, start_date: str):
    """
    Convert AI study plan JSON into list of DB-ready rows
    """
    items = []

    # ✅ Safe date handling
    if isinstance(start_date, str):
        try:
            base_date = datetime.fromisoformat(start_date).date()
        except Exception:
            base_date = datetime.today().date()
    elif isinstance(start_date, date):
        base_date = start_date
    else:
        base_date = datetime.today().date()

    # ✅ Extract correct key (VERY IMPORTANT FIX)
    days = ai_json.get("study_plan") or ai_json.get("days") or []

    print("DEBUG → Extracted days:", days)
    print("DEBUG → Days count:", len(days))

    if not isinstance(days, list) or len(days) == 0:
        print("❌ No valid study plan received from AI")
        return []

    # ✅ Normalize each day
    for idx, day in enumerate(days):
        print(f"Processing day {idx}: {day}")

        if not isinstance(day, dict):
            continue

        # ✅ Safe topics handling
        topics = day.get("topics", [])
        if isinstance(topics, list):
            topic_text = ", ".join(map(str, topics))
        else:
            topic_text = str(topics)

        # ✅ Safe tasks handling
        tasks = day.get("tasks", [])
        if isinstance(tasks, list):
            task_text = "\n".join(map(str, tasks))
        else:
            task_text = str(tasks)

        # ✅ Build item
        items.append({
            "date": base_date + timedelta(days=idx),  # ✅ FIXED
            "topic": topic_text,
            "hours": day.get("hours", 2),
            "tasks": task_text,
        })

    print("✅ Final items created:", len(items))

    return items

def parse_uploaded_file(file_or_path):
    """
    Wrapper used by tasks
    Returns:
    - normalized extracted text
    """

    text = extract_text_from_file(file_or_path)
    return normalize_text(text)
